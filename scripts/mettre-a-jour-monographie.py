from __future__ import annotations

import argparse
import os
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from PIL import Image, ImageDraw, ImageFont


BLEU = "#dbeafe"
VERT = "#dcfce7"
VIOLET = "#ede9fe"
ROSE = "#fce7f3"
JAUNE = "#fef3c7"
GRIS = "#f1f5f9"
TRAIT = "#334155"


def police(taille: int, gras: bool = False) -> ImageFont.FreeTypeFont:
    dossier = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    nom = "arialbd.ttf" if gras else "arial.ttf"
    return ImageFont.truetype(str(dossier / nom), taille)


def texte_centre(dessin: ImageDraw.ImageDraw, zone: tuple[int, int, int, int], texte: str, taille: int = 28, gras: bool = False) -> None:
    fonte = police(taille, gras)
    lignes = texte.split("\n")
    hauteurs = [dessin.textbbox((0, 0), ligne, font=fonte)[3] for ligne in lignes]
    y = zone[1] + (zone[3] - zone[1] - sum(hauteurs) - 8 * (len(lignes) - 1)) / 2
    for ligne, hauteur in zip(lignes, hauteurs):
        largeur = dessin.textbbox((0, 0), ligne, font=fonte)[2]
        dessin.text((zone[0] + (zone[2] - zone[0] - largeur) / 2, y), ligne, fill=TRAIT, font=fonte)
        y += hauteur + 8


def boite(dessin: ImageDraw.ImageDraw, zone: tuple[int, int, int, int], texte: str, couleur: str, taille: int = 27, rayon: int = 18) -> None:
    dessin.rounded_rectangle(zone, radius=rayon, fill=couleur, outline=TRAIT, width=3)
    texte_centre(dessin, zone, texte, taille, True)


def fleche(dessin: ImageDraw.ImageDraw, debut: tuple[int, int], fin: tuple[int, int], couleur: str = TRAIT, largeur: int = 3) -> None:
    dessin.line((debut, fin), fill=couleur, width=largeur)
    import math
    angle = math.atan2(fin[1] - debut[1], fin[0] - debut[0])
    longueur = 14
    for delta in (2.55, -2.55):
        point = (fin[0] + longueur * math.cos(angle + delta), fin[1] + longueur * math.sin(angle + delta))
        dessin.line((fin, point), fill=couleur, width=largeur)


def nouveau_fond(taille: tuple[int, int]) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", taille, "white")
    return image, ImageDraw.Draw(image)


def diagramme_contexte(chemin: Path) -> None:
    image, d = nouveau_fond((2048, 870))
    roles = ["SUPER_ADMINISTRATEUR", "ADMINISTRATEUR", "GESTIONNAIRE_BOURSES", "ÉVALUATEUR", "COMPTABLE", "ÉTUDIANT"]
    couleurs = [BLEU, BLEU, VERT, VIOLET, ROSE, JAUNE]
    systeme = (720, 310, 1400, 540)
    boite(d, systeme, "SYSTÈME DE GESTION DES BOURSES\nMINISTÈRE DE L’ESU", BLEU, 34)
    for index, (role, couleur) in enumerate(zip(roles, couleurs)):
        y = 35 + index * 132
        zone = (45, y, 455, y + 82)
        boite(d, zone, role, couleur, 23)
        fleche(d, (455, y + 41), (720, 350 + index * 27))
    services = [("MySQL / MariaDB\nAlwaysdata", (1640, 215, 1995, 365), VIOLET), ("Vercel Blob privé\nDocuments, preuves, logo", (1590, 515, 2010, 685), ROSE)]
    for texte, zone, couleur in services:
        boite(d, zone, texte, couleur, 25)
        fleche(d, (1400, 400 if "MySQL" in texte else 485), (zone[0], (zone[1] + zone[3]) // 2))
    d.text((790, 575), "Sessions JWT • API sécurisées • notifications • audit", fill=TRAIT, font=police(25, True))
    image.save(chemin)


def diagramme_cas_utilisation(chemin: Path) -> None:
    image, d = nouveau_fond((2048, 1230))
    d.rounded_rectangle((320, 35, 1760, 1185), radius=28, outline="#94a3b8", width=4, fill="#f8fafc")
    texte_centre(d, (420, 40, 1660, 120), "Système informatisé de gestion des bourses", 38, True)
    roles = [("Super administrateur", 120), ("Administrateur", 300), ("Gestionnaire", 480), ("Évaluateur", 660), ("Comptable", 840), ("Étudiant", 1020)]
    couleurs = [BLEU, BLEU, VERT, VIOLET, ROSE, JAUNE]
    cas = [
        ("Gérer comptes, rôles,\nparamètres et audit", 500, 155), ("Gérer établissements,\nprogrammes et appels", 1050, 155),
        ("S’inscrire et gérer\nson compte", 500, 355), ("Déposer candidature\net pièces privées", 1050, 355),
        ("Vérifier dossier et\npièces obligatoires", 500, 555), ("Affecter et évaluer\npar critères pondérés", 1050, 555),
        ("Formaliser décision et\ncréer l’attribution", 500, 755), ("Planifier et suivre\nles paiements", 1050, 755),
        ("Renouveler, suspendre,\nréactiver ou clôturer", 500, 955), ("Consulter notifications,\ntableau de bord et rapports", 1050, 955),
    ]
    liens = {0: [0, 1, 9], 1: [0, 1, 4, 6, 9], 2: [1, 4, 5, 6, 8, 9], 3: [5, 9], 4: [7, 9], 5: [2, 3, 9]}
    for role_index, cas_indices in liens.items():
        _, y = roles[role_index]
        for cas_index in cas_indices:
            _, x, cy = cas[cas_index]
            d.line(((300, y + 47), (x, cy + 62)), fill="#64748b", width=2)
    for (role, y), couleur in zip(roles, couleurs): boite(d, (30, y, 300, y + 95), role, couleur, 23)
    for texte, x, y in cas:
        zone = (x, y, x + 470, y + 125)
        d.ellipse(zone, fill="white", outline=TRAIT, width=3)
        texte_centre(d, zone, texte, 25)
    image.save(chemin)


def diagramme_workflow(chemin: Path) -> None:
    image, d = nouveau_fond((2048, 1230))
    colonnes = [("GESTIONNAIRE", BLEU), ("ÉTUDIANT", JAUNE), ("SYSTÈME", GRIS), ("ÉVALUATEUR / DÉCIDEUR", VIOLET), ("COMPTABLE", ROSE)]
    largeur = 390
    for index, (titre, couleur) in enumerate(colonnes):
        x = 45 + index * largeur
        d.rounded_rectangle((x, 35, x + largeur, 1180), radius=16, outline="#94a3b8", width=3)
        boite(d, (x, 35, x + largeur, 115), titre, couleur, 23)
    etapes = [
        (0, 160, "Configurer et publier\nl’appel"), (1, 275, "Créer le compte et\nla candidature brouillon"),
        (1, 425, "Téléverser les\npièces obligatoires"), (2, 565, "Contrôler la période,\nl’unicité et les pièces"),
        (0, 705, "Vérifier le dossier et\naffecter l’évaluateur"), (3, 835, "Noter les critères et\ncalculer le score pondéré"),
        (3, 985, "Formaliser la décision"), (2, 1085, "Créer attribution et\néchéancier si acceptée"),
    ]
    zones = []
    for col, y, texte in etapes:
        x = 75 + col * largeur
        zone = (x, y, x + 330, y + 95)
        boite(d, zone, texte, colonnes[col][1], 21)
        zones.append(zone)
    for depart, arrivee in zip(zones, zones[1:]):
        debut = ((depart[0] + depart[2]) // 2, depart[3])
        fin = ((arrivee[0] + arrivee[2]) // 2, arrivee[1])
        fleche(d, debut, fin)
    boite(d, (1635, 945, 1995, 1045), "Enregistrer paiements\net preuves", ROSE, 21)
    dernier = zones[-1]
    fleche(d, (dernier[2], (dernier[1] + dernier[3]) // 2), (1635, 995))
    d.text((1450, 1090), "Renouveler • suspendre • réactiver • clôturer", fill=TRAIT, font=police(20, True))
    image.save(chemin)


def diagramme_modele(chemin: Path, relationnel: bool = False) -> None:
    image, d = nouveau_fond((2048, 1400))
    titre = "SCHÉMA RELATIONNEL COMPLET" if relationnel else "DIAGRAMME DE CLASSES DU DOMAINE"
    texte_centre(d, (100, 15, 1948, 85), titre, 38, True)
    entites = [
        ("UTILISATEUR", "PK id\nUQ email\nrôle, actif\nsécurité compte", BLEU),
        ("ÉTUDIANT", "PK id\nUQ matricule\nidentité, formation\nFK établissement", VERT),
        ("ÉTABLISSEMENT", "PK id\nnom, sigle\nprovince, ville\nactif", VERT),
        ("PROGRAMME_BOURSE", "PK id\nUQ code\nfinancement, devise\ncibles", JAUNE),
        ("APPEL_CANDIDATURE", "PK id\nUQ référence\ndates, publication\nFK programme", JAUNE),
        ("DOCUMENT_REQUIS", "PK id\ntype, libellé\nobligatoire\nFK appel", ROSE),
        ("CRITÈRE_ÉVALUATION", "PK id\nlibellé, pondération\nnote maximale\nFK appel", VIOLET),
        ("CANDIDATURE", "PK id\nUQ référence\nUQ étudiant + appel\nstatut, score", VIOLET),
        ("DOCUMENT_CANDIDATURE", "PK id\ntype, fichier privé\nstatut vérification\nFK candidature", ROSE),
        ("HISTORIQUE_CANDIDATURE", "PK id\nancien/nouveau statut\ncommentaire\nFK candidature/utilisateur", GRIS),
        ("ÉVALUATION", "PK id\nUQ candidature + évaluateur\nnote, avis, échéance\nFK candidature/utilisateur", VIOLET),
        ("NOTE_CRITÈRE", "PK id\nUQ évaluation + critère\nnote, commentaire\nFK évaluation/critère", VIOLET),
        ("DÉCISION_CANDIDATURE", "PK id\ntype, justification\ndate décision\nFK candidature/décideur", JAUNE),
        ("ATTRIBUTION_BOURSE", "PK id\nUQ référence/candidature\nmontant, devise, dates\nFK étudiant", VERT),
        ("PAIEMENT", "PK id\nUQ référence\nmontant, échéance, preuve\nFK attribution", ROSE),
        ("RENOUVELLEMENT", "PK id\nUQ attribution + année\ndécision, suspension\nFK attribution", VERT),
        ("NOTIFICATION", "PK id\ntitre, message, lien\nlue\nFK utilisateur", BLEU),
        ("JOURNAL_AUDIT", "PK id\naction, entité, détails\ndate\nFK utilisateur", GRIS),
        ("PARAMÈTRE", "PK id\nUQ clé\nvaleur\ndates", GRIS),
    ]
    colonnes = 5
    largeur, hauteur = 360, 260
    positions: dict[str, tuple[int, int, int, int]] = {}
    for index, (nom, champs, couleur) in enumerate(entites):
        ligne, colonne = divmod(index, colonnes)
        x, y = 55 + colonne * 400, 105 + ligne * 315
        zone = (x, y, x + largeur, y + hauteur)
        positions[nom] = zone
        d.rounded_rectangle(zone, radius=14, fill="white", outline=TRAIT, width=3)
        d.rounded_rectangle((x, y, x + largeur, y + 58), radius=14, fill=couleur, outline=TRAIT, width=3)
        texte_centre(d, (x, y, x + largeur, y + 58), nom, 21, True)
        d.multiline_text((x + 18, y + 72), champs, fill=TRAIT, font=police(19), spacing=9)
    if relationnel:
        d.text((60, 1350), "PK : clé primaire     FK : clé étrangère     UQ : contrainte unique     Relations détaillées par les clés étrangères", fill=TRAIT, font=police(20, True))
    else:
        d.text((60, 1350), "Les associations traduisent le workflow : appel → candidature → évaluation/décision → attribution → paiements/renouvellements.", fill=TRAIT, font=police(20, True))
    image.save(chemin)


def remplacer_texte_paragraphe(paragraphe, texte: str) -> None:
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
        for run in paragraphe.runs[1:]: run.text = ""
    else:
        paragraphe.add_run(texte)


def remplacer_cellule(cellule, texte: str) -> None:
    paragraphe = cellule.paragraphs[0]
    remplacer_texte_paragraphe(paragraphe, texte)
    for autre in cellule.paragraphs[1:]: remplacer_texte_paragraphe(autre, "")


def actualiser_document(source: Path, destination: Path, dossier_diagrammes: Path) -> None:
    document = Document(source)
    remplacements = {
        "Évaluation et score : L’évaluateur saisit une note, une décision et un commentaire. Une contrainte unique évite plusieurs évaluations du même dossier par le même évaluateur. Le score final peut être recalculé à partir des évaluations enregistrées.": "Évaluation et score : Le gestionnaire configure les critères, leurs notes maximales et leurs pondérations pour chaque appel, puis affecte les évaluateurs. Chaque évaluateur ne traite que ses dossiers assignés. Les notes par critère produisent automatiquement un score pondéré sur 100 ; l’étudiant ne saisit jamais ce score.",
        "Attributions, paiements et renouvellements : Une attribution relie la candidature retenue à l’étudiant, au montant, à la devise et à la période. Les paiements enregistrent les échéances et statuts. Les renouvellements permettent de suivre la continuité de la bourse par année académique.": "Décisions, attributions et paiements : Une décision formelle acceptée ou rejetée conserve le décideur, la justification et son historique. L’acceptation crée l’attribution et son échéancier. Les paiements respectent le montant et la devise accordés, acceptent une preuve privée et ne peuvent dépasser l’attribution. Le suivi couvre renouvellement, suspension, réactivation, clôture et annulation.",
        "Rapports et impression : Les rapports proposent des filtres par période, établissement, programme, étudiant, statut et décision. Ils présentent des indicateurs, un tableau détaillé et une vue imprimable pouvant être enregistrée au format PDF.": "Pilotage, notifications et audit : Le tableau de bord présente des statistiques réelles, dossiers récents, échéances, alertes et paiements. Les rapports sont filtrables, exportables en CSV et imprimables en PDF. Les notifications informent les utilisateurs des événements importants et le journal d’audit retrace les actions sensibles.",
        "Déploiement : Le projet est connecté à GitHub et déployé sur Vercel. La base est distante sur Alwaysdata. Le Blob privé est connecté au projet Vercel. Les variables secrètes sont gérées dans l’environnement de déploiement.": "Déploiement : Le projet utilise Node.js 22, un build Next.js sur Vercel, MySQL/MariaDB sur Alwaysdata et Vercel Blob privé. Prisma 7 communique au moyen de l’adaptateur MariaDB et d’une configuration centralisée. DATABASE_URL, AUTH_SECRET, AUTH_URL et BLOB_READ_WRITE_TOKEN restent exclusivement dans les environnements autorisés.",
        "La solution obtenue couvre le cycle principal de gestion des bourses et ne se limite plus à des pages statiques. Les écrans sont reliés à des API et à la base de données. Les utilisateurs autorisés peuvent créer, modifier, consulter, filtrer et traiter les ressources. Les documents sont séparés des données structurées et protégés par un stockage privé. Les rapports consolident les informations enregistrées et facilitent l’impression.": "La solution obtenue couvre le cycle complet retenu pour le projet et ne contient plus de simulation métier. Les écrans utilisent les API et la base de données, les références sont générées côté serveur, les recherches sont automatiques et les autorisations dépendent du rôle. Les pièces, preuves de paiement et logos sont conservés dans un Blob privé. Les critères pondérés, décisions formelles, échéanciers, notifications, journaux d’audit et rapports assurent la continuité et la traçabilité du traitement.",
        "ajouter un moteur de critères pondérés configurable par programme ;": "ajouter l’authentification multifacteur pour les comptes administratifs ;",
        "mettre en place une journalisation d’audit détaillée ;": "mettre en place des tableaux de supervision et alertes de sécurité plus avancés ;",
        'DATABASE_URL="mysql://UTILISATEUR:MOT_DE_PASSE@HOTE:3306/BASE" AUTH_SECRET="CLE_SECRETE_LONGUE" AUTH_URL="http://localhost:3000" VERCEL_OIDC_TOKEN="fourni localement par Vercel" ADMIN_INITIAL_NOM="Administrateur ESU" ADMIN_INITIAL_EMAIL="adresse@domaine.cd" ADMIN_INITIAL_MOT_DE_PASSE="mot de passe fort"': 'DATABASE_URL="mysql://UTILISATEUR:MOT_DE_PASSE@HOTE:3306/BASE" AUTH_SECRET="CLE_SECRETE_LONGUE" AUTH_URL="https://application.vercel.app" BLOB_READ_WRITE_TOKEN="vercel_blob_rw_..." ADMIN_INITIAL_NOM="Administrateur ESU" ADMIN_INITIAL_EMAIL="adresse@domaine.cd" ADMIN_INITIAL_MOT_DE_PASSE="mot de passe fort (12 caractères minimum)"',
    }
    for paragraphe in document.paragraphs:
        if paragraphe.text in remplacements: remplacer_texte_paragraphe(paragraphe, remplacements[paragraphe.text])

    tables = document.tables
    mises_a_jour = {
        (5, 4, 0): "Tailwind CSS et Base UI", (5, 4, 2): "Interface responsive, accessible, cohérente et composants réutilisables.",
        (7, 9, 1): "Configurer les critères, affecter les évaluateurs et calculer automatiquement le score pondéré.",
        (7, 10, 1): "Formaliser une décision d’acceptation ou de rejet avec justification et historique.",
        (7, 15, 1): "Gérer paramètres, logo institutionnel, notifications et journal d’audit.",
        (9, 4, 1): "Une candidature ne peut être soumise que pendant l’ouverture et après dépôt de toutes les pièces obligatoires.",
        (9, 5, 1): "Un évaluateur ne peut évaluer qu’une fois un dossier qui lui a été affecté.",
        (10, 4, 1): "Code généré, financement, plafond, devise, cibles et activation.",
        (10, 5, 1): "Référence générée, calendrier, places, pièces requises, critères pondérés et publication.",
        (10, 8, 1): "Affectations, notes par critère, avis et score pondéré automatique.",
        (10, 9, 0): "Décisions et attributions", (10, 9, 1): "Décision formelle, justification, avis imprimable, montant, devise, période et échéancier automatique.",
        (10, 14, 0): "Paramètres, notifications et audit", (10, 14, 1): "Identité visuelle, centre de notifications et traçabilité des actions sensibles.",
        (11, 7, 1): "Créer, déposer les pièces, soumettre, vérifier, affecter, évaluer et décider", (11, 7, 2): "Conforme",
        (11, 8, 1): "Créer automatiquement attribution et échéancier, puis contrôler montant, devise et preuve", (11, 8, 2): "Conforme",
        (13, 2, 2): "Ajout de @prisma/adapter-mariadb et conversion centralisée de DATABASE_URL en configuration du pilote, supprimant le délai d’attente du pool.",
        (14, 7, 2): "id, type, fichier privé, statut et motif de vérification, candidatureId, vérificateurId.",
        (14, 8, 2): "id, statut, échéance, note calculée, avis, candidatureId, evaluateurId.",
        (15, 5, 5): "-", (15, 5, 6): "C/R/U/D (propres pièces)",
        (16, 7, 2): "Afficher la candidature, les pièces requises, leur consultation et leur vérification.",
        (16, 8, 2): "Afficher l’affectation, les critères pondérés, le score et la décision formelle.",
    }
    for (table, ligne, colonne), valeur in mises_a_jour.items(): remplacer_cellule(tables[table].rows[ligne].cells[colonne], valeur)

    dictionnaire = tables[14]
    entites_manquantes = [
        ("DocumentRequisAppel", "Pièce attendue pour un appel.", "id, type, libellé, estObligatoire, appelId."),
        ("HistoriqueCandidature", "Trace des transitions du dossier.", "id, ancienStatut, nouveauStatut, commentaire, candidatureId, utilisateurId."),
        ("CritereEvaluation", "Critère pondéré défini pour un appel.", "id, libellé, pondération, noteMaximale, appelId."),
        ("NoteCritereEvaluation", "Note attribuée à un critère.", "id, note, commentaire, evaluationId, critereId."),
        ("DecisionCandidature", "Décision formelle et justifiée.", "id, type, justification, dateDecision, candidatureId, decideurId."),
        ("JournalAudit", "Trace d’une action sensible.", "id, action, entité, entiteId, détails, utilisateurId, date."),
    ]
    noms = {row.cells[0].text for row in dictionnaire.rows}
    for valeurs in entites_manquantes:
        if valeurs[0] in noms: continue
        cellules = dictionnaire.add_row().cells
        for cellule, valeur in zip(cellules, valeurs): remplacer_cellule(cellule, valeur)

    document.save(destination)

    diagramme_contexte(dossier_diagrammes / "image1.png")
    diagramme_cas_utilisation(dossier_diagrammes / "image2.png")
    diagramme_workflow(dossier_diagrammes / "image3.png")
    diagramme_modele(dossier_diagrammes / "image7.png", False)
    diagramme_modele(dossier_diagrammes / "image8.png", True)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=destination.parent) as temporaire:
        chemin_temporaire = Path(temporaire.name)
    try:
        with ZipFile(destination, "r") as source_zip, ZipFile(chemin_temporaire, "w", ZIP_DEFLATED) as destination_zip:
            remplaces = {f"word/media/image{numero}.png": dossier_diagrammes / f"image{numero}.png" for numero in (1, 2, 3, 7, 8)}
            for info in source_zip.infolist():
                contenu = remplaces[info.filename].read_bytes() if info.filename in remplaces else source_zip.read(info.filename)
                destination_zip.writestr(info, contenu)
        chemin_temporaire.replace(destination)
    finally:
        if chemin_temporaire.exists(): chemin_temporaire.unlink()


def main() -> None:
    analyseur = argparse.ArgumentParser(description="Actualise la monographie BESU sans modifier le document original.")
    analyseur.add_argument("source", type=Path)
    analyseur.add_argument("destination", type=Path)
    analyseur.add_argument("--diagrammes", type=Path, default=Path(".analyse-rapport/diagrammes-mis-a-jour"))
    arguments = analyseur.parse_args()
    arguments.diagrammes.mkdir(parents=True, exist_ok=True)
    actualiser_document(arguments.source, arguments.destination, arguments.diagrammes)
    print(arguments.destination)


if __name__ == "__main__":
    main()
